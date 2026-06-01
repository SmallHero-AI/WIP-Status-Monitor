import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

document = Document()

# Set extremely tight margins to guarantee 1 page
sections = document.sections
for section in sections:
    section.top_margin = Cm(0.8)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

# Add Title
p = document.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('高雄市稅捐稽徵處 115 年度公教員工生活津貼申請表暨收據')
run.font.name = '標楷體'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
run.font.size = Pt(14)
run.bold = True

# Create Table (17 rows, 16 cols)
table = document.add_table(rows=17, cols=16)
table.style = 'Table Grid'
table.autofit = False
table.allow_autofit = False

# Force equal column widths (21 - 1.6 = 19.4 cm / 16 = 1.2125 cm)
for row in table.rows:
    for cell in row.cells:
        cell.width = Cm(1.21)

def merge_cells(r, c_start, c_end):
    if c_start == c_end: return table.cell(r, c_start)
    cell = table.cell(r, c_start)
    cell.merge(table.cell(r, c_end))
    return cell

def merge_range(r_start, r_end, c_start, c_end):
    cell = table.cell(r_start, c_start)
    cell.merge(table.cell(r_end, c_end))
    return cell

def set_text(cell, text):
    cell.text = text

# Row 0
set_text(merge_cells(0, 0, 1), '申請人姓名')
set_text(merge_cells(0, 2, 3), '')
set_text(merge_cells(0, 4, 5), '單    位')
set_text(merge_cells(0, 6, 8), '')
set_text(merge_range(0, 1, 9, 11), '申  請\n日  期')
set_text(merge_range(0, 1, 12, 15), '    年    月    日')

# Row 1
set_text(merge_cells(1, 0, 1), '員 工 代 號')
set_text(merge_cells(1, 2, 3), '')
set_text(merge_cells(1, 4, 5), '職    稱')
set_text(merge_cells(1, 6, 8), '')

# Row 2
set_text(merge_cells(2, 0, 1), '補 助 事 由')
set_text(merge_cells(2, 2, 15), '□ 一、結婚補助\n□ 二、生育補助\n□ 三、喪葬補助\n□ 四、子女教育補助')

# Row 3
set_text(merge_cells(3, 0, 1), '檢 附 證 件')
set_text(merge_cells(3, 2, 15), '□ 一、結婚證明文件、戶口名簿(各二份)\n□ 二、出生證明書、戶口名簿(各二份)\n□ 三、死亡證明書、戶口名簿(各二份)\n□ 四、收費單據 (國中、國小無須繳驗)、戶口名簿 (於本機關第一次申請時須繳驗)')

# Row 4
set_text(merge_range(4, 5, 0, 1), '婚、喪、生\n育  補  助')
set_text(merge_cells(4, 2, 3), '事實發生日期')
set_text(merge_cells(4, 4, 6), '    年  月  日')
set_text(merge_cells(4, 7, 9), '關係人身分證字號')
set_text(merge_cells(4, 10, 10), '')
set_text(merge_cells(4, 11, 13), '關係人姓名')
set_text(merge_cells(4, 14, 14), '稱謂')
set_text(merge_cells(4, 15, 15), '')

# Row 5
set_text(merge_cells(5, 2, 3), '薪(俸)額')
set_text(merge_cells(5, 4, 6), '')
set_text(merge_cells(5, 7, 9), '補助月俸數')
set_text(merge_cells(5, 10, 12), '□五個月 □三個月 □二個月')
set_text(merge_cells(5, 13, 14), '申請補助金額')
set_text(merge_cells(5, 15, 15), '')

# Row 6
set_text(merge_range(6, 7, 0, 1), '子        女\n教 育 補 助')
set_text(merge_cells(6, 2, 3), '子 女 姓 名')
set_text(merge_cells(6, 4, 5), '')
set_text(merge_cells(6, 6, 7), '就 讀 學 校')
set_text(merge_cells(6, 8, 9), '')
set_text(merge_cells(6, 10, 10), '年級')
set_text(merge_cells(6, 11, 11), '')
set_text(merge_cells(6, 12, 13), '就讀學校支給標準')
set_text(merge_cells(6, 14, 15), '申請補助金額')

# Row 7
set_text(merge_cells(7, 2, 3), '身分證字號')
set_text(merge_cells(7, 4, 5), '')
set_text(merge_cells(7, 6, 11), '')
set_text(merge_cells(7, 12, 13), '□公立   □私立\n□夜間部 □其他')
set_text(merge_cells(7, 14, 15), '')

# Row 8
set_text(merge_range(8, 9, 0, 0), '區分')
set_text(merge_cells(8, 1, 3), '大學暨獨立學院')
set_text(merge_cells(8, 4, 6), '五專後二年、二、三專')
set_text(merge_cells(8, 7, 8), '五專前三年')
set_text(merge_cells(8, 9, 10), '高中')
set_text(merge_cells(8, 11, 13), '高職')
set_text(merge_cells(8, 14, 14), '國中')
set_text(merge_cells(8, 15, 15), '國小')

# Row 9
set_text(table.cell(9, 1), '公立')
set_text(table.cell(9, 2), '私立')
set_text(table.cell(9, 3), '夜間部')
set_text(table.cell(9, 4), '公立')
set_text(table.cell(9, 5), '私立')
set_text(table.cell(9, 6), '夜間部')
set_text(table.cell(9, 7), '公立')
set_text(table.cell(9, 8), '私立')
set_text(table.cell(9, 9), '公立')
set_text(table.cell(9, 10), '私立')
set_text(table.cell(9, 11), '公立')
set_text(table.cell(9, 12), '私立')
set_text(table.cell(9, 13), '實用技能班')
set_text(table.cell(9, 14), '公立')
set_text(table.cell(9, 15), '公立')

# Row 10
set_text(table.cell(10, 0), '支給數額')
vals = ['13600', '35800', '14300', '10000', '28000', '14300', '7700', '20800', '3800', '13500', '3200', '18900', '1500', '500', '500']
for i, v in enumerate(vals):
    set_text(table.cell(10, i+1), v)

# Row 11
set_text(merge_cells(11, 0, 1), '核 准 補 助\n金        額')
set_text(merge_cells(11, 2, 15), '新臺幣          拾          萬          仟          佰          拾          元整')

# Row 12
set_text(merge_cells(12, 0, 1), '領        款\n收        據')
set_text(merge_cells(12, 2, 4), '下列金額業已\n如數收訖此據')
set_text(merge_cells(12, 5, 10), '新臺幣   拾   萬   仟   佰   拾   元整')
set_text(merge_cells(12, 11, 12), '具領人\n簽  章')
set_text(merge_cells(12, 13, 15), '')

# Row 13
set_text(merge_cells(13, 0, 1), '入 匯 帳 號')
set_text(merge_cells(13, 2, 8), '高雄銀行               分行')
set_text(merge_cells(13, 9, 15), '(請填高雄銀行之薪資帳號)')

# Row 14
set_text(merge_cells(14, 0, 1), '切\n\n結\n\n書')
text_terms = '''一、本表所列各項補助，依「全國軍公教員工待遇支給要點」之其他給與各附表規定辦理，其所附證件皆屬事實，如有虛報、冒領、重領或偽造、變造、虛偽欺瞞等情事，除繳回所領金額外，並願接受一切行政及法令處分。
二、子女以未婚且無職業需仰賴申請人扶養者為限。公教人員申請子女教育補助時，其未婚子女如繼續從事經常性工作，且註冊之日前六個月工作平均每月所得（依所得稅法申報之所得）超過申請時之勞工基本工資者，以有職業論，不得申請補助。
三、公教人員子女具有下列情形之一者，不得申請子女教育補助。但不包括領取優秀學生獎學金、清寒獎學金、民間團體獎學金及就讀國中小未因特殊身分獲有全免(減免)學雜費或政府提供獎助者：
(一)全免或減免學雜費(含十二年國民基本教育學費補助)。(二)屬未具學籍之學校或補習班學生。(三)就讀公私立中等以上學校之選讀生。(四)就讀無特定修業年限之學校。(五)已獲有軍公教遺族就學費用優待條例享有公費、減免學雜費之優待。(六)已領取其他政府提供之獎(補)助。
四、公教人員子女除就讀國中小未因特殊身分全免(減免)學雜費及政府提供獎助者，依表訂數額申請子女教育補助外，其實際繳納之學雜費低於子女教育補助表訂數額者，僅得申請補助其實際繳納數額。
五、公教人員請領子女教育補助，應以在職期間其子女已完成當學期註冊手續為要件。其申請以各級學校所規定之修業年限為準。如有轉學、轉系、重考、留級、重修情形，其於同一學制重複就讀之年級，不再補助。又畢業後再考入相同學制學校就讀者，不得請領。
六、夫妻同為公教人員者，其子女教育補助應自行協調由一方申領。
七、死亡之父母、配偶以未擔任公職者為限；夫妻或其他親屬同為公教人員者，對同一死亡事實，以報領一份喪葬補助為限。
                                                                                                    簽章：'''
set_text(merge_cells(14, 2, 15), text_terms)

# Row 15 & 16
set_text(merge_cells(15, 0, 3), '人 事 室')
set_text(merge_cells(15, 4, 7), '秘 書 室')
set_text(merge_cells(15, 8, 11), '會 計 室')
set_text(merge_range(15, 16, 12, 15), '\n\n機 關 首 長')

set_text(merge_cells(16, 0, 3), '\n\n\n')
set_text(merge_cells(16, 4, 7), '\n\n\n')
set_text(merge_cells(16, 8, 11), '\n\n\n')

# Format alignment, spacing, and fonts
for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            if '本表所列各項補助' not in p.text and '一、結婚補助' not in p.text and '結婚證明文件' not in p.text and '□五個月' not in p.text and '□公立' not in p.text:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for run in p.runs:
                run.font.name = '標楷體'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                if i >= 8 and i <= 10:
                    run.font.size = Pt(8) # smaller for grid
                elif '本表所列各項補助' in p.text:
                    run.font.size = Pt(8.5) # very compact for terms
                else:
                    run.font.size = Pt(9.5)

out_path = r'e:\G-AI-1\高雄市稅捐稽徵處115年度公教員工生活津貼申請表暨收據.docx'
document.save(out_path)
print(f"Document saved to {out_path}")
