import urllib.request
import urllib.parse
from docx import Document

url = "https://www.kh.edu.tw/files/bulletin/161476/114%E5%B9%B4%E5%BA%A6%E5%85%AC%E6%95%99%E5%93%A1%E5%B7%A5%E7%94%9F%E6%B4%BB%E6%B4%A5%E8%B2%BC%E7%94%B3%E8%AB%8B%E8%A1%A8%E6%9A%A8%E6%94%B6%E6%93%9A.docx"
local_filename = r"e:\G-AI-1\temp_official_form.docx"
final_filename = r"e:\G-AI-1\高雄市稅捐稽徵處115年度公教員工生活津貼申請表暨收據_官方下載版.docx"

try:
    print("Downloading...")
    # Add headers to avoid 403 Forbidden
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(local_filename, 'wb') as out_file:
        data = response.read()
        out_file.write(data)
    print("Download complete.")
    
    # Try to open the document and replace the title
    try:
        doc = Document(local_filename)
        # Search for the title paragraph and replace it
        for p in doc.paragraphs:
            if "公教員工生活津貼申請表暨收據" in p.text:
                print(f"Found title: {p.text}")
                # Clear all runs and replace with new title
                for run in p.runs:
                    run.text = ""
                p.add_run("高雄市稅捐稽徵處 115 年度公教員工生活津貼申請表暨收據")
        
        # Save to the final filename
        doc.save(final_filename)
        print(f"Successfully modified and saved to {final_filename}")
    except Exception as e:
        print(f"Error modifying document: {e}")
        # If we can't parse it with python-docx, just save the raw file
        import shutil
        shutil.copy2(local_filename, final_filename)
        print(f"Saved raw file to {final_filename} instead.")

except Exception as e:
    print(f"Error downloading: {e}")
